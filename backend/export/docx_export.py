"""DOCX export generator for conversations using python-docx."""

from typing import Dict, Any, List, Optional
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

from .content_processing import (
    get_display_name,
    process_export_content,
    build_extended_character_names
)
from .table_parser import parse_markdown_tables
import re


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph, converting markdown to DOCX formatting.

    Converts:
    - ### Header -> bold text (header level becomes just bold)
    - **text** -> bold
    - *text* or _text_ -> italic
    - `text` -> monospace
    """
    if not text:
        return ""

    # First, convert markdown headers (# ## ###) to bold at start of line
    # This simplifies the text before inline processing
    lines = text.split('\n')
    processed_lines = []
    for line in lines:
        line = line.strip()
        # Match markdown headers at start of line
        if line.startswith('#'):
            # Count the # symbols
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                content = match.group(2).strip()
                # Convert to just bold (we don't change paragraph style)
                processed_lines.append(f"**{content}**")
            else:
                processed_lines.append(line)
        else:
            processed_lines.append(line)

    text = '\n'.join(processed_lines)

    # Regex pattern for markdown formatting
    # Matches **bold**, *italic*, _italic_, and `code`
    # Note: \*[^*]+?\* matches single asterisk italic (not **bold**)
    pattern = r'(\*\*.*?\*\*|_.*?_|\*[^*]+?\*|`.*?`)'

    parts = re.split(pattern, text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if not part:
            i += 1
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            # Italic with underscores
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic with asterisks (only if not bold)
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code/monospace
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            # Plain text - preserve line breaks
            if '\n' in part:
                # Split by newlines and add each as separate run
                subparts = part.split('\n')
                for j, subpart in enumerate(subparts):
                    if subpart:
                        paragraph.add_run(subpart)
                    if j < len(subparts) - 1:
                        # Add line break within paragraph
                        paragraph.add_run().add_break()
            else:
                paragraph.add_run(part)
        i += 1


def _setup_document_styles(doc: Document) -> None:
    """Set up document styles for the export."""
    # Modify Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def _add_content_to_docx(doc: Document, text: str) -> None:
    """Add content to DOCX, rendering GFM pipe tables as native table objects and inline markdown."""
    segments = parse_markdown_tables(text)
    for seg in segments:
        if seg["type"] == "text":
            content = seg["content"]
            if content.strip():
                # Create paragraph and add formatted text
                p = doc.add_paragraph()
                _add_formatted_text(p, content)
        else:
            headers = seg["headers"]
            rows = seg["rows"]
            if not headers:
                continue

            num_cols = len(headers)
            tbl = doc.add_table(rows=1, cols=num_cols)
            tbl.style = 'Table Grid'

            # Header row -- bold
            header_cells = tbl.rows[0].cells
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Data rows
            for row_data in rows:
                row_cells = tbl.add_row().cells
                for i, cell_text in enumerate(row_data):
                    # Guard against rows with fewer cells than headers
                    if i < len(row_cells):
                        row_cells[i].text = cell_text

            doc.add_paragraph()  # spacing after table


def _add_stage1_docx(
    doc: Document,
    stage1: List[Dict[str, Any]],
    character_names: Dict[str, str]
) -> None:
    """Add Stage 1 responses to the document."""
    doc.add_heading("Stage 1: Initial Responses", level=2)

    for idx, response in enumerate(stage1):
        model_id = response.get("model", "Unknown")
        instance_key = f"{model_id}:{idx}"
        display_name = get_display_name(model_id, character_names, instance_key)

        doc.add_heading(display_name, level=3)

        content = response.get("response", "")
        error = response.get("error")

        if error:
            p = doc.add_paragraph()
            run = p.add_run(f"Error: {error}")
            run.italic = True
        elif content:
            _add_content_to_docx(doc, content)


def _add_stage2_docx(
    doc: Document,
    stage2: List[Dict[str, Any]],
    character_names: Dict[str, str],
    aggregate_rankings: Optional[List[Dict[str, Any]]] = None,
    label_to_model: Optional[Dict[str, str]] = None,
    label_to_instance_key: Optional[Dict[str, str]] = None
) -> None:
    """Add Stage 2 rankings to the document.

    Args:
        doc: The Document object
        stage2: List of ranking responses
        character_names: Dict mapping instance keys to character names
        aggregate_rankings: Optional aggregate rankings data
        label_to_model: Dict mapping anonymous labels to model IDs
        label_to_instance_key: Dict mapping anonymous labels to instance keys
    """
    doc.add_heading("Stage 2: Peer Rankings", level=2)

    for idx, ranking in enumerate(stage2):
        model_id = ranking.get("model", "Unknown")
        instance_key = f"{model_id}:{idx}"
        display_name = get_display_name(model_id, character_names, instance_key)

        doc.add_heading(f"{display_name}'s Ranking", level=3)

        content = ranking.get("ranking", "")
        error = ranking.get("error")

        if error:
            p = doc.add_paragraph()
            run = p.add_run(f"Error: {error}")
            run.italic = True
        elif content:
            # De-anonymize the ranking content
            processed_content = process_export_content(
                content,
                label_to_model=label_to_model,
                label_to_instance_key=label_to_instance_key,
                character_names=character_names
            )
            _add_content_to_docx(doc, processed_content)

    # Add aggregate rankings table if available
    if aggregate_rankings:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Aggregate Rankings:")
        run.bold = True

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Rank"
        header_cells[1].text = "Model"
        header_cells[2].text = "Average Score"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for idx, item in enumerate(aggregate_rankings, 1):
            model_id = item.get("model", "Unknown")
            instance_key = item.get("instance_key", "")
            avg_rank = item.get("average_rank", 0)

            display_name = get_display_name(model_id, character_names, instance_key)

            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = display_name
            row_cells[2].text = f"{avg_rank:.1f}"


def _add_stage3_docx(
    doc: Document,
    stage3: List[Dict[str, Any]],
    character_names: Dict[str, str]
) -> None:
    """Add Stage 3 revisions to the document."""
    doc.add_heading("Stage 3: Revisions", level=2)

    for idx, response in enumerate(stage3):
        model_id = response.get("model", "Unknown")
        instance_key = f"{model_id}:{idx}"
        display_name = get_display_name(model_id, character_names, instance_key)

        doc.add_heading(display_name, level=3)

        content = response.get("response", "")
        error = response.get("error")

        if error:
            p = doc.add_paragraph()
            run = p.add_run(f"Error: {error}")
            run.italic = True
        elif content:
            _add_content_to_docx(doc, content)


def _add_stage4_docx(
    doc: Document,
    stage4: Dict[str, Any],
    character_names: Dict[str, str]
) -> None:
    """Add Stage 4 analysis (truth-check, rankings, highlights) to the document."""
    doc.add_heading("Stage 4: Analysis", level=2)

    # --- Truth-Check ---
    doc.add_heading("Truth-Check Results", level=3)
    truth_check = stage4.get("truth_check") or {}
    if truth_check.get("checked") and truth_check.get("claims"):
        summary = truth_check.get("summary", {})
        verified = summary.get("confirmed", 0)
        flagged = summary.get("disputed", 0)
        unverifiable = summary.get("unaddressed", 0)
        doc.add_paragraph(
            f"Summary: {verified} confirmed, {flagged} disputed, {unverifiable} unaddressed"
        )
        for claim in truth_check["claims"]:
            verdict = claim.get("verdict", "Unknown")
            text = claim.get("text", "")

            # New fields from quick-052 and quick-054
            priority = claim.get("priority", "medium")
            searched = claim.get("searched", True)
            reason = claim.get("reason", "")

            # Build claim line with priority badge
            priority_badge = f"[{priority.upper()}]" if priority != "medium" else ""
            searched_note = " (not verified)" if not searched else ""

            claim_parts = []
            if priority_badge:
                claim_parts.append(priority_badge)
            claim_parts.append(f"[{verdict}]")
            claim_parts.append(text)  # Keep raw markdown for formatting
            if searched_note:
                claim_parts.append(searched_note)

            claim_line = " ".join(claim_parts)
            # Create paragraph with formatted text (markdown will be converted)
            p = doc.add_paragraph()
            _add_formatted_text(p, claim_line)

            # Add reason for all verdicts if provided
            if reason:
                p_reason = doc.add_paragraph()
                p_reason.add_run("  - Reason: ")
                _add_formatted_text(p_reason, reason)
    else:
        doc.add_paragraph("Truth-check was not run for this conversation.")

    # --- Rankings ---
    doc.add_heading("Chairman Rankings", level=3)
    rankings_data = stage4.get("rankings") or {}
    rankings_list = rankings_data.get("rankings", [])
    if rankings_data.get("checked") and rankings_list:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(['Rank', 'Model', 'Score', 'Reasoning', 'Insight', 'Clarity']):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for rank_num, entry in enumerate(rankings_list, 1):
            display_name = get_display_name(entry.get("model", "Unknown"), character_names)
            score = str(int(entry.get("normalized_score", 0)))
            dims = entry.get("dimensions", {})
            reasoning = dims.get("reasoning", {}).get("label", "")
            insight = dims.get("insight", {}).get("label", "")
            clarity = dims.get("clarity", {}).get("label", "")

            row_cells = table.add_row().cells
            row_cells[0].text = str(rank_num)
            row_cells[1].text = display_name
            row_cells[2].text = score
            row_cells[3].text = reasoning
            row_cells[4].text = insight
            row_cells[5].text = clarity
    else:
        doc.add_paragraph("Rankings were not produced for this conversation.")

    # --- Highlights ---
    doc.add_heading("Council Highlights", level=3)
    highlights = stage4.get("highlights") or {}
    agreements = highlights.get("agreements", [])
    disagreements = highlights.get("disagreements", [])
    unique_insights = highlights.get("unique_insights", [])

    if agreements or disagreements or unique_insights:
        if agreements:
            p = doc.add_paragraph()
            p.add_run("Agreements:").bold = True
            for item in agreements:
                model_names = ", ".join(
                    get_display_name(m, character_names) for m in item.get("models", [])
                )
                finding = item.get('finding', '')
                # Create paragraph with formatted markdown
                p = doc.add_paragraph()
                _add_formatted_text(p, f"- {finding} (models: {model_names})")
        if disagreements:
            p = doc.add_paragraph()
            p.add_run("Disagreements:").bold = True
            for item in disagreements:
                topic = item.get("topic", "")
                why = item.get("why_they_differ", "")
                positions = item.get("positions", [])
                p = doc.add_paragraph()
                p.add_run(topic).bold = True
                if why:
                    p2 = doc.add_paragraph()
                    _add_formatted_text(p2, f"  Why they differ: {why}")
                for pos in positions:
                    name = get_display_name(pos.get("model_id", "Unknown"), character_names)
                    p3 = doc.add_paragraph()
                    _add_formatted_text(p3, f"  {name}: {pos.get('position_text', '')}")
        if unique_insights:
            p = doc.add_paragraph()
            p.add_run("Unique Insights:").bold = True
            for item in unique_insights:
                model_name = get_display_name(item.get("model", "Unknown"), character_names)
                finding = item.get('finding', '')
                # Create paragraph with formatted markdown
                p = doc.add_paragraph()
                _add_formatted_text(p, f"- {finding} ({model_name})")
    else:
        doc.add_paragraph("No highlights were extracted for this conversation.")

    # --- Debates ---
    debates = stage4.get("debates") or []
    completed_debates = [d for d in debates if d.get("transcript")]
    if completed_debates:
        doc.add_heading("Debates", level=3)
        for debate in completed_debates:
            title = debate.get("title", "Debate")
            doc.add_heading(title, level=4)

            participants = debate.get("participants", [])
            if participants:
                parts_str = ", ".join(
                    f"{p.get('name', 'Unknown')} ({p.get('role', '')})"
                    for p in participants
                )
                p = doc.add_paragraph()
                p.add_run("Participants: ").bold = True
                p.add_run(parts_str)

            transcript = debate.get("transcript", [])
            if transcript:
                p = doc.add_paragraph()
                p.add_run("Transcript").bold = True
                for turn in transcript:
                    name = turn.get("name", turn.get("role", "Unknown"))
                    text = turn.get("text", "")
                    p = doc.add_paragraph()
                    p.add_run(f"{name}: ").bold = True
                    _add_formatted_text(p, text)

            verdict = (debate.get("verdict") or {}).get("summary", "")
            if verdict:
                p = doc.add_paragraph()
                p.add_run("Verdict: ").bold = True
                _add_formatted_text(p, verdict)


def _add_stage5_docx(
    doc: Document,
    stage5: Dict[str, Any],
    character_names: Dict[str, str],
    chairman_character_name: Optional[str] = None
) -> None:
    """Add Stage 5 synthesis to the document."""
    doc.add_heading("Stage 5: Chairman Synthesis", level=2)

    model_id = stage5.get("model", "Unknown")
    display_name = chairman_character_name if chairman_character_name else get_display_name(model_id, character_names)

    doc.add_heading(display_name, level=3)

    content = stage5.get("response", "")
    error = stage5.get("error")

    if error:
        p = doc.add_paragraph()
        run = p.add_run(f"Error: {error}")
        run.italic = True
    elif content:
        _add_content_to_docx(doc, content)


def _add_metadata_docx(
    doc: Document,
    conversation: Dict[str, Any],
    council_config: Optional[Dict[str, Any]] = None
) -> None:
    """Add metadata section to the document."""
    # Separator
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("─" * 50)

    doc.add_heading("Conversation Details", level=2)

    # Basic info
    created_at = conversation.get("created_at", "Unknown")
    doc.add_paragraph(f"Timestamp: {created_at}")

    # Find execution mode
    execution_mode = "Unknown"
    for msg in conversation.get("messages", []):
        if msg.get("role") == "assistant" and msg.get("metadata"):
            execution_mode = msg["metadata"].get("execution_mode", "Unknown")
            break

    doc.add_paragraph(f"Execution Mode: {execution_mode}")

    # Council config info
    if council_config:
        council_models = council_config.get("council_models", [])
        chairman_model = council_config.get("chairman_model", "")
        character_names = council_config.get("character_names", {}) or {}
        chairman_character_name = council_config.get("chairman_character_name")

        # Member list
        member_names = []
        for idx, model_id in enumerate(council_models):
            instance_key = f"{model_id}:{idx}"
            name = character_names.get(instance_key) or character_names.get(str(idx)) or get_display_name(model_id, {})
            member_names.append(name)

        doc.add_paragraph(f"Council Members: {', '.join(member_names) if member_names else 'None'}")

        chairman_display = chairman_character_name if chairman_character_name else get_display_name(chairman_model, {})
        doc.add_paragraph(f"Chairman: {chairman_display}")

        # Configuration
        doc.add_heading("Council Configuration", level=3)
        if council_config.get("council_temperature") is not None:
            doc.add_paragraph(f"Council Temperature: {council_config['council_temperature']}")
        if council_config.get("chairman_temperature") is not None:
            doc.add_paragraph(f"Chairman Temperature: {council_config['chairman_temperature']}")
        if council_config.get("stage2_temperature") is not None:
            doc.add_paragraph(f"Stage 2 Temperature: {council_config['stage2_temperature']}")
        if council_config.get("revision_temperature") is not None:
            doc.add_paragraph(f"Revision Temperature: {council_config['revision_temperature']}")

        # System prompts (truncated)
        doc.add_heading("System Prompts", level=3)
        max_prompt_len = 500

        for stage_name, display_name in [
            ("stage1", "Stage 1"),
            ("stage2", "Stage 2"),
            ("stage5", "Stage 5"),
            ("revision", "Revision")
        ]:
            prompt = council_config.get(f"{stage_name}_prompt")
            if prompt:
                p = doc.add_paragraph()
                run = p.add_run(f"{display_name} Prompt:")
                run.bold = True

                if len(prompt) > max_prompt_len:
                    prompt = prompt[:max_prompt_len] + "..."
                doc.add_paragraph(prompt)

    # Web search info
    for msg in conversation.get("messages", []):
        if msg.get("role") == "assistant" and msg.get("metadata"):
            metadata = msg["metadata"]
            search_query = metadata.get("search_query")

            if search_query:
                doc.add_heading("Web Search", level=3)
                doc.add_paragraph(f"Query: {search_query}")
            break


def export_docx(conversation: Dict[str, Any]) -> bytes:
    """
    Export a conversation as a DOCX document.

    Args:
        conversation: The conversation dict to export

    Returns:
        DOCX bytes
    """
    doc = Document()
    _setup_document_styles(doc)

    # Get council config
    council_config = conversation.get("council_config", {})
    character_names = build_extended_character_names(council_config)
    chairman_character_name = council_config.get("chairman_character_name")

    # Iterate through messages
    for message in conversation.get("messages", []):
        role = message.get("role")

        if role == "user":
            content = message.get("content", "")
            doc.add_heading(content, level=1)

        elif role == "assistant":
            # Stage 1 - always present
            stage1 = message.get("stage1", [])
            if stage1:
                _add_stage1_docx(doc, stage1, character_names)

            # Stage 2 - peer rankings
            stage2 = message.get("stage2", [])
            if stage2:
                metadata = message.get("metadata", {})
                aggregate_rankings = metadata.get("aggregate_rankings") if metadata else None
                label_to_model = metadata.get("label_to_model") if metadata else None
                label_to_instance_key = metadata.get("label_to_instance_key") if metadata else None
                _add_stage2_docx(
                    doc,
                    stage2,
                    character_names,
                    aggregate_rankings,
                    label_to_model,
                    label_to_instance_key
                )

            # Stage 3 - revisions
            stage3 = message.get("stage3")
            if stage3:
                _add_stage3_docx(doc, stage3, character_names)

            # Stage 4 - analysis
            stage4 = message.get("stage4")
            if stage4:
                _add_stage4_docx(doc, stage4, character_names)

            # Stage 5 - chairman synthesis
            stage5 = message.get("stage5")
            if stage5:
                _add_stage5_docx(doc, stage5, character_names, chairman_character_name)

    # Metadata section at the end
    _add_metadata_docx(doc, conversation, council_config)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()
